from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIRECTORY = PROJECT_ROOT / "knowledge-base"


def set_font(run, size, bold=False, color="000000"):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def create_internal_guidelines_docx():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.25

    noise = document.add_paragraph()
    noise.alignment = WD_ALIGN_PARAGRAPH.CENTER
    noise.paragraph_format.space_after = Pt(4)
    set_font(noise.add_run("Seite 1"), 9, color="777777")

    separator = document.add_paragraph()
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator.paragraph_format.space_after = Pt(10)
    set_font(separator.add_run("----------------------------------------"), 9, color="999999")

    company = document.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company.paragraph_format.space_after = Pt(4)
    set_font(company.add_run("AlexZ Job Assistant"), 10, True, "2E74B5")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(14)
    set_font(title.add_run("Interne Richtlinie zur Dokumentenverarbeitung"), 20, True, "1F4D78")

    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    set_font(
        intro.add_run(
            "Diese interne Richtlinie beschreibt, wie Mitarbeiter Dokumente im AlexZ Job Assistant prüfen und verarbeiten."
        ),
        11,
    )

    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(6)
    set_font(heading.add_run("Prüfung vor dem Upload"), 14, True, "2E74B5")

    for item in (
        "Nur Dateien im Format TXT, PDF oder DOCX verwenden.",
        "Vor dem Upload prüfen, ob das Dokument lesbaren Text enthält.",
        "Dateien mit Passwörtern, API-Schlüsseln oder unnötigen personenbezogenen Daten ablehnen.",
    ):
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        set_font(paragraph.add_run(item), 11)

    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(6)
    set_font(heading.add_run("Kontrolle nach dem Upload"), 14, True, "2E74B5")

    control = document.add_paragraph()
    set_font(
        control.add_run(
            "Nach der Verarbeitung kontrolliert der Mitarbeiter die Anzahl der erzeugten Chunks in der Antwort und die neuen Points in Qdrant. Der gespeicherte Chunk muss bereinigten Inhalt enthalten. Seitenangaben, Trennlinien und Copyright-Zeilen dürfen nicht gespeichert werden."
        ),
        11,
    )

    repeated_company = document.add_paragraph()
    repeated_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    repeated_company.paragraph_format.space_before = Pt(14)
    repeated_company.paragraph_format.space_after = Pt(3)
    set_font(repeated_company.add_run("AlexZ Job Assistant"), 9, True, "777777")

    copyright_line = document.add_paragraph()
    copyright_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        copyright_line.add_run("© AlexZ Automation 2026 - Nur für interne Verwendung"),
        8,
        color="777777",
    )

    document.save(OUTPUT_DIRECTORY / "internal-guidelines.docx")


def create_public_faq_pdf():
    pdfmetrics.registerFont(
        TTFont(
            "DejaVuSans",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        )
    )
    pdfmetrics.registerFont(
        TTFont(
            "DejaVuSans-Bold",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        )
    )

    output_path = OUTPUT_DIRECTORY / "public-faq.pdf"
    document = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title="AlexZ Job Assistant - Öffentliche FAQ",
        author="AlexZ Automation",
    )

    styles = getSampleStyleSheet()
    noise_style = ParagraphStyle(
        "Noise",
        parent=styles["Normal"],
        fontName="DejaVuSans",
        fontSize=8,
        textColor=HexColor("#777777"),
        alignment=TA_CENTER,
        leading=10,
        spaceAfter=4,
    )
    company_style = ParagraphStyle(
        "Company",
        parent=noise_style,
        fontSize=10,
        textColor=HexColor("#2E74B5"),
        leading=12,
        spaceAfter=6,
    )
    title_style = ParagraphStyle(
        "FaqTitle",
        parent=styles["Title"],
        fontName="DejaVuSans-Bold",
        fontSize=22,
        textColor=HexColor("#1F4D78"),
        alignment=TA_CENTER,
        leading=26,
        spaceAfter=18,
    )
    question_style = ParagraphStyle(
        "Question",
        parent=styles["Heading2"],
        fontName="DejaVuSans-Bold",
        fontSize=12,
        textColor=HexColor("#2E74B5"),
        leading=15,
        spaceBefore=10,
        spaceAfter=4,
    )
    answer_style = ParagraphStyle(
        "Answer",
        parent=styles["BodyText"],
        fontName="DejaVuSans",
        fontSize=10.5,
        leading=15,
        spaceAfter=8,
    )

    story = [
        Paragraph("Seite 1", noise_style),
        Paragraph("----------------------------------------", noise_style),
        Paragraph("AlexZ Job Assistant", company_style),
        Paragraph("Häufig gestellte Fragen", title_style),
        Paragraph("Welche Dokumentformate werden unterstützt?", question_style),
        Paragraph("Der Upload unterstützt Textdateien im TXT-Format, PDF-Dokumente mit auslesbarem Text und Word-Dokumente im DOCX-Format.", answer_style),
        Paragraph("Was geschieht nach dem Upload?", question_style),
        Paragraph("Der Text wird extrahiert, von technischen Elementen bereinigt, in Chunks aufgeteilt und als Vektoren in der Wissensbasis gespeichert.", answer_style),
        Paragraph("Werden gescannte PDF-Dateien erkannt?", question_style),
        Paragraph("Nein. In dieser Version wird kein OCR verwendet. Eine PDF-Datei muss bereits auslesbaren Text enthalten.", answer_style),
        Paragraph("Was passiert bei einem falschen Dateiformat?", question_style),
        Paragraph("Das System beendet die Verarbeitung und gibt eine verständliche Fehlermeldung zurück. Unterstützt werden ausschließlich TXT, PDF und DOCX.", answer_style),
        Spacer(1, 16),
        Paragraph("AlexZ Job Assistant", noise_style),
        Paragraph("© AlexZ Automation 2026", noise_style),
    ]

    document.build(story)


if __name__ == "__main__":
    OUTPUT_DIRECTORY.mkdir(parents=True, exist_ok=True)
    create_internal_guidelines_docx()
    create_public_faq_pdf()
